"""
DOC/DOCX Text Extraction Pipeline
Extracts text content from Microsoft Word documents
"""

import os
import logging
from typing import List, Dict, Optional
from docx import Document

# Optional magic import for MIME type detection
try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False

logger = logging.getLogger(__name__)

class DOCExtractor:
    """Extract text content from DOC and DOCX files"""
    
    def __init__(self):
        self.supported_extensions = ['.doc', '.docx']
        self.mime_types = [
            'application/msword',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        ]
    
    def is_supported(self, file_path: str) -> bool:
        """Check if the file is a supported DOC/DOCX"""
        if not os.path.exists(file_path):
            return False
        
        # Check file extension
        _, ext = os.path.splitext(file_path.lower())
        if ext not in self.supported_extensions:
            return False
        
        # Check MIME type if magic is available
        if MAGIC_AVAILABLE:
            try:
                mime_type = magic.from_file(file_path, mime=True)
                return mime_type in self.mime_types
            except Exception as e:
                logger.warning(f"Could not determine MIME type for {file_path}: {e}")
                return True  # Assume it's valid if we can't check MIME type
        else:
            # Fallback: just check file extension
            return True
    
    def extract_text(self, file_path: str) -> Dict[str, any]:
        """
        Extract text from a DOC/DOCX file
        
        Args:
            file_path: Path to the DOC/DOCX file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        if not self.is_supported(file_path):
            raise ValueError(f"File {file_path} is not a supported DOC/DOCX")
        
        try:
            logger.info(f"Extracting text from DOC/DOCX: {file_path}")
            
            # Read DOCX file
            doc = Document(file_path)
            
            # Extract metadata
            metadata = {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'file_size': os.path.getsize(file_path),
                'num_paragraphs': len(doc.paragraphs),
                'num_tables': len(doc.tables),
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'subject': doc.core_properties.subject or '',
                'keywords': doc.core_properties.keywords or '',
                'comments': doc.core_properties.comments or '',
                'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                'modified': str(doc.core_properties.modified) if doc.core_properties.modified else '',
                'last_modified_by': doc.core_properties.last_modified_by or '',
                'revision': doc.core_properties.revision or 0,
                'version': doc.core_properties.version or ''
            }
            
            # Extract text from paragraphs
            paragraph_texts = []
            full_text = ""
            
            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                para_text = paragraph.text.strip()
                if para_text:  # Only include non-empty paragraphs
                    paragraph_texts.append({
                        'paragraph_number': para_num,
                        'text': para_text,
                        'char_count': len(para_text),
                        'style': paragraph.style.name if paragraph.style else 'Normal'
                    })
                    full_text += para_text + "\n"
            
            # Extract text from tables
            table_texts = []
            for table_num, table in enumerate(doc.tables, 1):
                table_text = ""
                table_data = []
                
                for row_num, row in enumerate(table.rows, 1):
                    row_data = []
                    for cell_num, cell in enumerate(row.cells, 1):
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                        table_text += cell_text + "\t"
                    table_data.append(row_data)
                    table_text += "\n"
                
                if table_text.strip():
                    table_texts.append({
                        'table_number': table_num,
                        'text': table_text.strip(),
                        'char_count': len(table_text.strip()),
                        'rows': len(table.rows),
                        'columns': len(table.columns) if table.rows else 0,
                        'data': table_data
                    })
                    full_text += table_text + "\n"
            
            result = {
                'success': True,
                'metadata': metadata,
                'full_text': full_text.strip(),
                'paragraph_texts': paragraph_texts,
                'table_texts': table_texts,
                'total_characters': len(full_text.strip()),
                'extraction_method': 'python-docx'
            }
            
            logger.info(f"Successfully extracted {len(full_text)} characters from {metadata['num_paragraphs']} paragraphs and {metadata['num_tables']} tables")
            return result
            
        except Exception as e:
            logger.error(f"Error extracting text from DOC/DOCX {file_path}: {e}")
            return {
                'success': False,
                'error': str(e),
                'file_path': file_path,
                'file_name': os.path.basename(file_path)
            }
    
    def extract_from_directory(self, directory_path: str) -> List[Dict[str, any]]:
        """
        Extract text from all DOC/DOCX files in a directory
        
        Args:
            directory_path: Path to directory containing DOC/DOCX files
            
        Returns:
            List of extraction results for each DOC/DOCX file
        """
        if not os.path.exists(directory_path):
            raise ValueError(f"Directory {directory_path} does not exist")
        
        results = []
        doc_files = [f for f in os.listdir(directory_path) 
                    if f.lower().endswith(('.doc', '.docx'))]
        
        logger.info(f"Found {len(doc_files)} DOC/DOCX files in {directory_path}")
        
        for doc_file in doc_files:
            file_path = os.path.join(directory_path, doc_file)
            try:
                result = self.extract_text(file_path)
                results.append(result)
            except Exception as e:
                logger.error(f"Failed to process {doc_file}: {e}")
                results.append({
                    'success': False,
                    'error': str(e),
                    'file_path': file_path,
                    'file_name': doc_file
                })
        
        return results
